import sys, os
sys.stdout.reconfigure(encoding='utf-8')

base = os.path.join(os.path.dirname(__file__), '..', 'references', 'แบบฟอร์ม')
out = os.path.join(os.path.dirname(__file__), 'forms_output.txt')

with open(out, 'w', encoding='utf-8') as f:

    # 1. Read PDF
    import fitz
    pdf_path = os.path.join(base, 'ตัวอย่าง2-1 2-2.pdf')
    doc = fitz.open(pdf_path)
    f.write("=" * 60 + "\n")
    f.write("  FILE: ตัวอย่าง2-1 2-2.pdf\n")
    f.write("=" * 60 + "\n")
    for i, page in enumerate(doc):
        f.write(f"\n--- Page {i+1} ---\n")
        f.write(page.get_text())
    doc.close()

    # 2. Read Excel files
    try:
        import openpyxl
    except ImportError:
        os.system('pip install openpyxl')
        import openpyxl

    for xls_name in ['บัญชี 1-1.xlsx', 'บัญชี 1-2.xlsx', '2แบบฟอร์มตารางรายงานผล 130168 กนต.xlsx']:
        xls_path = os.path.join(base, xls_name)
        if os.path.exists(xls_path):
            f.write("\n\n" + "=" * 60 + "\n")
            f.write(f"  FILE: {xls_name}\n")
            f.write("=" * 60 + "\n")
            wb = openpyxl.load_workbook(xls_path)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                f.write(f"\n  Sheet: {sheet_name} ({ws.max_row} rows x {ws.max_column} cols)\n")
                f.write("-" * 40 + "\n")
                for row in ws.iter_rows(max_row=min(ws.max_row, 30), values_only=False):
                    vals = [str(cell.value) if cell.value is not None else '' for cell in row]
                    f.write(" | ".join(vals) + "\n")

    # 3. Read Word files
    try:
        import docx
    except ImportError:
        os.system('pip install python-docx')
        import docx

    for doc_name in ['แบบฟอร์ม 2-1.doc', 'แบบฟอร์ม 2-2.doc', 'แบบฟอร์ม 2-3.doc', 'หนังสือรับรองตนเอง ปรับแก้เหลือ 1 หน้า.docx']:
        doc_path = os.path.join(base, doc_name)
        if os.path.exists(doc_path):
            f.write("\n\n" + "=" * 60 + "\n")
            f.write(f"  FILE: {doc_name}\n")
            f.write("=" * 60 + "\n")
            if doc_name.endswith('.docx'):
                try:
                    d = docx.Document(doc_path)
                    for para in d.paragraphs:
                        if para.text.strip():
                            f.write(para.text + "\n")
                    for table in d.tables:
                        f.write("\n[TABLE]\n")
                        for row in table.rows:
                            vals = [cell.text.strip() for cell in row.cells]
                            f.write(" | ".join(vals) + "\n")
                except Exception as e:
                    f.write(f"Error: {e}\n")
            else:
                f.write("(.doc format - cannot read directly, need antiword or libreoffice)\n")

print(f"Done! Output -> {out}")
